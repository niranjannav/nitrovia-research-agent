"""DOCX renderer using python-docx."""

import logging
from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models.schemas import GeneratedReport, ReportSection

logger = logging.getLogger(__name__)


class DOCXRenderer:
    """Renders reports as DOCX documents."""

    # Brand colors
    PRIMARY_COLOR = RGBColor(37, 99, 235)  # #2563eb
    SECONDARY_COLOR = RGBColor(30, 64, 175)  # #1e40af
    TEXT_COLOR = RGBColor(51, 51, 51)  # #333333

    def render(self, report: GeneratedReport) -> BytesIO:
        """
        Render a report to DOCX.

        Args:
            report: Generated report content

        Returns:
            BytesIO containing DOCX data
        """
        try:
            doc = Document()

            # Set up document styles
            self._setup_styles(doc)

            # Title
            title = doc.add_heading(report.title, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.color.rgb = self.PRIMARY_COLOR

            doc.add_paragraph()  # Spacer

            # Executive Summary
            doc.add_heading("Executive Summary", 1)
            summary_para = doc.add_paragraph(report.executive_summary)
            summary_para.style = "Normal"

            # Main Sections
            for section in report.sections:
                self._add_section(doc, section, level=1)

            # Key Findings
            doc.add_heading("Key Findings", 1)
            for finding in report.key_findings:
                para = doc.add_paragraph(finding, style="List Bullet")

            # Recommendations
            doc.add_heading("Recommendations", 1)
            for rec in report.recommendations:
                para = doc.add_paragraph(rec, style="List Bullet")

            # Sources
            if report.sources:
                doc.add_heading("Sources", 1)
                for source in report.sources:
                    para = doc.add_paragraph(source, style="List Bullet")
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(107, 114, 128)

            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            return output

        except Exception as e:
            logger.error(f"DOCX rendering failed: {e}")
            raise

    def _setup_styles(self, doc: Document) -> None:
        """Configure document styles."""
        styles = doc.styles

        # Modify Normal style
        normal_style = styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = self.TEXT_COLOR
        normal_style.paragraph_format.space_after = Pt(8)
        normal_style.paragraph_format.line_spacing = 1.15

        # Modify Heading 1
        h1_style = styles["Heading 1"]
        h1_style.font.name = "Calibri"
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = self.PRIMARY_COLOR
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(6)

        # Modify Heading 2
        h2_style = styles["Heading 2"]
        h2_style.font.name = "Calibri"
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = self.SECONDARY_COLOR
        h2_style.paragraph_format.space_before = Pt(14)
        h2_style.paragraph_format.space_after = Pt(4)

        # Modify Heading 3
        h3_style = styles["Heading 3"]
        h3_style.font.name = "Calibri"
        h3_style.font.size = Pt(12)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(55, 65, 81)
        h3_style.paragraph_format.space_before = Pt(10)
        h3_style.paragraph_format.space_after = Pt(4)

    def _add_section(
        self,
        doc: Document,
        section: ReportSection,
        level: int,
    ) -> None:
        """
        Add a section to the document recursively.

        Args:
            doc: Document instance
            section: Section to add
            level: Heading level (1-3)
        """
        # Add heading (cap at level 3)
        heading_level = min(level, 3)
        doc.add_heading(section.title, heading_level)

        # Add content
        if section.content:
            # Split content into paragraphs
            paragraphs = section.content.split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.style = "Normal"

        # Add subsections recursively
        for subsection in section.subsections:
            self._add_section(doc, subsection, level + 1)
